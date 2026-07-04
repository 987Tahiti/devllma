"""
Lecture/ecriture de documents bureautiques (Word, Excel, PDF) pour l'agent generaliste.
Permet a l'agent d'agir sur de vrais documents (pas seulement du texte/code brut),
comme demande : "il doit pouvoir agir sur des documents, les lire et faire des
modifications sur le document".
"""
import os
import docx
from openpyxl import Workbook, load_workbook
from pypdf import PdfReader
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas

DOC_EXTS = {".docx", ".xlsx", ".pdf"}

def is_office_doc(path):
    return os.path.splitext(path)[1].lower() in DOC_EXTS

def read_document(path):
    """Extrait le texte d'un document Word/Excel/PDF. Retourne une chaine texte."""
    ext = os.path.splitext(path)[1].lower()
    if ext == ".docx":
        d = docx.Document(path)
        parts = [p.text for p in d.paragraphs]
        for table in d.tables:
            for row in table.rows:
                parts.append(" | ".join(c.text for c in row.cells))
        return "\n".join(parts)
    if ext == ".xlsx":
        wb = load_workbook(path, data_only=True)
        parts = []
        for sheet in wb.worksheets:
            parts.append(f"### Feuille: {sheet.title}")
            for row in sheet.iter_rows(values_only=True):
                parts.append("\t".join("" if c is None else str(c) for c in row))
        return "\n".join(parts)
    if ext == ".pdf":
        reader = PdfReader(path)
        parts = [f"### Page {i+1}\n{(page.extract_text() or '').strip()}"
                 for i, page in enumerate(reader.pages)]
        return "\n\n".join(parts)
    raise ValueError(f"format de document non pris en charge: {ext}")

def write_document(path, content):
    """(Re)cree un document Word/Excel/PDF a partir de texte.
    Word: un paragraphe par ligne. Excel: '### Feuille: X' delimite les feuilles,
    les lignes suivantes sont separees par tabulation ou virgule -> une ligne de cellules.
    PDF: genere un nouveau PDF texte simple (pas d'edition fine d'un PDF existant,
    la structure interne d'un PDF ne s'y prete pas de maniere fiable)."""
    ext = os.path.splitext(path)[1].lower()
    os.makedirs(os.path.dirname(path) or ".", exist_ok=True)
    if ext == ".docx":
        d = docx.Document()
        for line in content.split("\n"):
            d.add_paragraph(line)
        d.save(path)
        return
    if ext == ".xlsx":
        wb = Workbook()
        wb.remove(wb.active)
        sheet = None
        for line in content.split("\n"):
            if line.startswith("### Feuille:"):
                sheet = wb.create_sheet(line.split(":", 1)[1].strip()[:31])
                continue
            if sheet is None:
                sheet = wb.create_sheet("Feuille1")
            cells = line.split("\t") if "\t" in line else line.split(",")
            sheet.append(cells)
        if not wb.worksheets:
            wb.create_sheet("Feuille1")
        wb.save(path)
        return
    if ext == ".pdf":
        c = canvas.Canvas(path, pagesize=A4)
        width, height = A4
        y = height - 50
        for line in content.split("\n"):
            if y < 50:
                c.showPage()
                y = height - 50
            c.drawString(40, y, line[:110])
            y -= 14
        c.save()
        return
    raise ValueError(f"format de document non pris en charge: {ext}")
